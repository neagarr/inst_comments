# docx_scraper.py
from docx import Document
import sys
from db_model import Session, CommentAnalysis


session = Session()

# Словарь: commenter_account_id → id последнего комментария этого пользователя
last_comment_by_user = {}


def get_hyperlink_info(hyperlink_block, doc):
    """
    Получаем текст и ссылку из параграфа с <w:hyperlink>
    """
    hyperlink_text = ''
    hyperlink_url = ''

    for node in hyperlink_block.iter():
        # текст внутри <w:t>
        if node.tag.endswith('}t') and node.text:
            hyperlink_text += node.text
        # ищем <w:hyperlink r:id="...">
        if 'hyperlink' in node.tag.lower():
            r_id = node.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            if r_id and r_id in doc.part.rels:
                hyperlink_url = doc.part.rels[r_id].target_ref
    return hyperlink_text.strip(), hyperlink_url


def extract_summary_before_first_image(docx_path, start_marker):
    from docx import Document

    doc = Document(docx_path)
    collecting = False
    collected_text = []

    for block in doc.element.body:
        tag = block.tag.lower()

        if 'p' in tag:
            paragraph_text = ''.join(
                node.text for node in block.iter() if node.tag.endswith('}t') and node.text
            )

            # Начало сбора
            if not collecting and start_marker in paragraph_text.lower().replace(' ', ''):
                collecting = True
                split_text = paragraph_text.split('=', 1)
                if len(split_text) > 1:
                    first_line = split_text[1].strip()
                    if first_line:
                        collected_text.append(first_line)
                continue

            if collecting:
                if paragraph_text.strip():
                    collected_text.append(paragraph_text.strip())

                # Проверка на картинку внутри параграфа
                has_drawing = any('drawing' in node.tag.lower() for node in block.iter())
                if has_drawing:
                    break

    return '\n\n'.join(collected_text)


def extract_comments(docx_path):
    doc = Document(docx_path)
    comments = []

    blocks = list(doc.element.body)
    collecting_summary = False
    i = 0
    total_blocks = len(blocks)

    while i < total_blocks:
        block = blocks[i]
        tag = block.tag.lower()

        # ищем начало после client_post_summary
        if 'p' in tag and not collecting_summary:
            paragraph_text = ''.join(
                node.text for node in block.iter() if node.tag.endswith('}t') and node.text
            )
            if 'client_post_summary=' in paragraph_text.lower().replace(' ', ''):
                collecting_summary = True
            i += 1
            continue

        if not collecting_summary:
            i += 1
            continue

        # если нашли картинку
        has_drawing = any('drawing' in node.tag.lower() for node in block.iter())

        if has_drawing:
            # определяем сколько картинок подряд
            start_i = i
            image_chain = 1
            while i + image_chain < total_blocks:
                next_block = blocks[i + image_chain]
                if any('drawing' in n.tag.lower() for n in next_block.iter()):
                    image_chain += 1
                else:
                    break

            # проверяем, идёт ли за последней картинкой гиперссылка
            last_img_idx = i + image_chain - 1
            next_block_after_images = (
                blocks[last_img_idx + 1] if last_img_idx + 1 < total_blocks else None
            )

            # если после цепочки картинок идёт hyperlink → последняя — аватар нового комментатора
            is_next_avatar = (
                next_block_after_images is not None
                and any('hyperlink' in n.tag.lower() for n in next_block_after_images.iter())
            )

            if is_next_avatar:
                # значит, после этой цепочки начинается новый комментарий
                comment = {'commenter_account_id': '', 'commenter_account_link': '', 'comment_text': ''}

                account_id, account_link = get_hyperlink_info(next_block_after_images, doc)
                comment['commenter_account_id'] = account_id
                comment['commenter_account_link'] = account_link

                i = last_img_idx + 2  # переходим после гиперссылки
                comment_text = []

                # собираем текст до следующей цепочки картинок
                while i < total_blocks:
                    inner_block = blocks[i]
                    has_inner_img = any('drawing' in n.tag.lower() for n in inner_block.iter())

                    # если следующая цепочка картинок → проверим, не аватар ли она
                    if has_inner_img:
                        # проверяем, идёт ли за этой цепочкой hyperlink — тогда выходим
                        j = i
                        while j < total_blocks and any(
                            'drawing' in n.tag.lower() for n in blocks[j].iter()
                        ):
                            j += 1
                        if j < total_blocks and any(
                            'hyperlink' in n.tag.lower() for n in blocks[j].iter()
                        ):
                            break  # конец текущего комментария
                    else:
                        text = ''.join(
                            n.text for n in inner_block.iter() if n.tag.endswith('}t') and n.text
                        ).strip()
                        if text:
                            comment_text.append(text)

                    i += 1

                comment['comment_text'] = '\n'.join(comment_text)
                comments.append(comment)
                continue

        i += 1

    return comments

def parse_docx(path):
    doc = Document(path)
    text_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # 1️⃣ client_account_id — первая строка
    client_account_id = text_lines[0].replace("client_account_id =", "").strip()

    # 2️⃣ client_account_link — вторая строка
    client_account_link = text_lines[1].replace("client_account_link =", "").strip()

    # 3️⃣ client_post_id — тр строка
    client_post_id = text_lines[2].replace("client_post_id =", "").strip()

    # 4️⃣ client_post_summary
    client_post_summary = extract_summary_before_first_image(path, "client_post_summary=")

    # 5️⃣ comments
    comments = extract_comments(path)

    # 6️⃣ сохраняем в базу
    for comment in comments:
        commenter_id = comment['commenter_account_id']
        comment_text = comment['comment_text'].strip()
        parent_id = None

        # 🔹 если комментарий начинается с @ — ищем родителя
        if comment_text.startswith("@"):
            mentioned_user = comment_text.split()[0][1:].strip().rstrip('.,:;!?')
            if mentioned_user in last_comment_by_user:
                parent_id = last_comment_by_user[mentioned_user]

        # 🔹 создаем запись
        record = CommentAnalysis(
            comment_text=comment_text,
            commenter_account_id=commenter_id,
            commenter_account_link=comment['commenter_account_link'],
            client_account_id=client_account_id,
            client_account_link=client_account_link,
            client_post_id=client_post_id,
            client_post_summary=client_post_summary.strip(),
            parent_comment_id=parent_id
        )
        session.add(record)
        session.commit()  # сохраняем, чтобы получить ID

        # 🔹 обновляем "последний комментарий пользователя"
        last_comment_by_user[commenter_id] = record.id

    print(f"✅ Импортировано {len(comments)} комментариев из {path}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("❗ Укажи путь к .docx файлу, например:")
        print("   python docx_scraper.py post_1.docx")
        sys.exit(1)

    file_path = sys.argv[1]
    parse_docx(file_path)

# for i in {1..9}; do python docx_scraper.py post_$i.docx; done
